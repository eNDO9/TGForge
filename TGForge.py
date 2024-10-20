import streamlit as st
from telethon.sync import TelegramClient
from telethon import TelegramClient as AsyncTelegramClient, functions
from telethon.errors import SessionPasswordNeededError
import nest_asyncio
import asyncio
import pandas as pd
from docx import Document
import os

# Apply nest_asyncio to allow nested event loops
nest_asyncio.apply()

# Load default credentials from st.secrets
default_api_id = st.secrets["telegram"].get("api_id", "")
default_api_hash = st.secrets["telegram"].get("api_hash", "")
default_phone = st.secrets["telegram"].get("phone", "")

# Step 1: Create Streamlit input fields for user credentials
st.title("Telegram Authentication")
st.write("Enter your Telegram API credentials or use the saved defaults.")

# Prompt user to input their credentials, with defaults pre-filled
api_id = st.text_input("API ID", value=default_api_id)
api_hash = st.text_input("API Hash", value=default_api_hash)
phone = st.text_input("Phone Number (e.g., +1 5718671248)", value=default_phone)

# Check to ensure values are being captured correctly
st.write(f"API ID: {api_id}, API Hash: {api_hash}, Phone: {phone}")

client = None
async_client = None

# Define explicit session file path, ensure directory is visible
session_dir = os.getcwd()  # Current directory, you can change to a specific path if needed
session_path = os.path.join(session_dir, "my_telegram_session")

# Function to delete the session file
def delete_session_file(session_path):
    session_file = f"{session_path}.session"
    if os.path.exists(session_file):
        os.remove(session_file)
        st.write("Existing session file deleted. You can try authenticating again.")

# Automatically check and clean if necessary
if "retry_authentication" in st.session_state and st.session_state["retry_authentication"]:
    delete_session_file(session_path)
    st.session_state["retry_authentication"] = False

# Check if session file already exists, indicating a previous successful login
if os.path.exists(f"{session_path}.session"):
    st.session_state['authenticated'] = True
else:
    st.session_state['authenticated'] = False

if api_id and api_hash and phone:
    try:
        # Use a specified session path in the local directory
        client = TelegramClient(session_path, api_id, api_hash)
        
        # Check if the session is already authenticated
        if st.session_state.get('authenticated'):
            st.success("Already authenticated. Skipping reauthentication.")
            
            # Ensure async_client is initialized even when reusing the session
            async_client = AsyncTelegramClient(session_path, api_id, api_hash)
        else:
            st.write("Credentials loaded. You can proceed with authentication.")
    except Exception as e:
        if "database is locked" in str(e).lower():
            st.error("Database is locked. Trying to resolve...")
            # Set a flag to delete the session file and retry authentication
            st.session_state["retry_authentication"] = True
            st.experimental_rerun()
        else:
            st.error(f"Error initializing Telegram client: {e}")

# Function to authenticate synchronously
def authenticate_client():
    global client, async_client
    try:
        client.connect()
        if not client.is_user_authorized():
            client.send_code_request(phone)
            st.write("A verification code has been sent to your Telegram account.")
            
            verification_code = st.text_input("Enter the verification code:", "")
            if st.button("Verify"):
                client.sign_in(phone, verification_code)
                st.success("Authentication successful!")

                # Create an async client for further use, with the same session path
                async_client = AsyncTelegramClient(session_path, api_id, api_hash)
                st.session_state['authenticated'] = True

        else:
            # Create async client if already authorized
            async_client = AsyncTelegramClient(session_path, api_id, api_hash)
            st.success("Already authenticated. Async client ready for further operations.")
            st.session_state['authenticated'] = True

    except SessionPasswordNeededError:
        st.error("Your account is protected by a password. Please disable it for this demo.")
    except Exception as e:
        if "database is locked" in str(e).lower():
            st.error("Database is locked. Attempting to delete the session file and retry...")
            # Trigger session file deletion and retry
            st.session_state["retry_authentication"] = True
            st.experimental_rerun()
        else:
            st.error(f"Error during authentication: {e}")
    finally:
        client.disconnect()

# Function to provide download link for the session file
def provide_session_download(session_path):
    session_file = f"{session_path}.session"
    if os.path.exists(session_file):
        with open(session_file, "rb") as file:
            btn = st.download_button(
                label="Download Session File",
                data=file,
                file_name="my_telegram_session.session",
                mime="application/octet-stream"
            )
            if btn:
                st.write("Session file downloaded. Save it securely.")

# Step 2: UI to switch to the new feature once authenticated
if st.session_state['authenticated']:
    # Hide the authentication fields and show new functionality
    st.title("Fetch Telegram Channel Information")
    
    # Input: channel names
    channel_list_input = st.text_input(
        "Please enter the channel name(s). Separate multiple channels with a comma:",
        placeholder="e.g., channel_1, channel_2"
    )
    channel_list = [channel.strip() for channel in channel_list_input.split(",")]
    
    # Select output format
    output_choice = st.selectbox(
        "How would you like to save the output?",
        ('xlsx', 'docx', 'print')
    )

    # Create a Word document object
    doc = Document()
    doc.add_heading('Telegram Channel Information', level=1)

    async def get_first_valid_message_date(channel_name):
        """Finds the date of the earliest available user-generated message in a channel."""
        try:
            channel = await async_client.get_entity(channel_name)
            async for message in async_client.iter_messages(channel, reverse=True):
                if message and not message.action:
                    if message.text or message.media:
                        return message.date.isoformat()
            return "No user-generated messages found"
        except Exception as e:
            st.write(f"Error fetching the first message for {channel_name}: {e}")
            return "Not Available"

    async def get_channel_info(channel_name):
        """Fetches detailed information about a Telegram channel."""
        try:
            result = await async_client(functions.channels.GetFullChannelRequest(channel=channel_name))

            first_message_date = await get_first_valid_message_date(channel_name)
            chat = result.chats[0]

            # Extract relevant channel information
            title = chat.title
            description = result.full_chat.about.strip() if result.full_chat.about else 'No Description'
            participants_count = result.full_chat.participants_count if hasattr(result.full_chat, 'participants_count') else 'Not Available'

            # Display or return the gathered information as required
            st.write(f"Channel Information for {title}:")
            channel_info = {
                "Title": title,
                "Description": description,
                "Participants": participants_count,
                "First Message Date": first_message_date
            }

            if output_choice == 'docx':
                doc.add_heading(f"Channel: {title}", level=2)
                for key, value in channel_info.items():
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(f"{key}: ").bold = True
                    paragraph.add_run(str(value))
                doc.add_paragraph("="*40)

            return channel_info

        except Exception as e:
            st.write(f"Error fetching info for {channel_name}: {e}")
            return None

    # Button to start gathering information
    if st.button("Fetch Information"):
        asyncio.run(get_channel_info(channel_list[0]))
        
        # Provide output options
        if output_choice == 'xlsx':
            st.write("Saving as Excel (not yet implemented).")
        elif output_choice == 'docx':
            doc.save('Telegram_Channel_Info.docx')
            st.download_button(
                label="Download Document",
                data=open('Telegram_Channel_Info.docx', 'rb').read(),
                file_name='Telegram_Channel_Info.docx'
            )
        elif output_choice == 'print':
            st.write("Printing data to the screen (not yet implemented).")

else:
    if st.button("Authenticate"):
        if client:
            authenticate_client()
        else:
            st.error("Please provide valid API ID, API Hash, and Phone Number.")